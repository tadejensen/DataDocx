from typing import Union, Optional

import docx.shape
import docx.shared
import docx.table
import database_functions as dbf
import gui_functions as gui_f
import pandas as pd
import time
from lengthy_imports import *
import pickle
from ipydex import IPS, activate_ips_on_exception
import os
import shutil
import copy
import subprocess
import hashlib as hl

import docx
from docx.shared import Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

config = load_config()


activate_ips_on_exception()

idx = pd.IndexSlice
#
# TODO: make entire program database-based
# keep backups for every project in project's folder
# (databases with few entries)
# (backup subsubfolder or whatever)
# also backup checklist in project's folder
# maybe backup current remarks_db in projects folder?

class Inspection():
    def __init__(self,
                 kind: str,
                 has_happened: bool,
                 parent_report,
                 scope: str='Prüfung der gesamten Anlage ohne Rotorblätter',
                 wea_hours: Optional[int] = None,
                 wea_output: Optional[int] = None,
                 inspectors_list: Optional[list] = None,
                 testrun_wind: Optional[float] = None,
                 testrun_power: Optional[float] = None,
                 timeline: Optional[str] = '',
                 ):
        self.kind = kind.upper()
        self.scope=scope
        self.has_happened = has_happened
        self.parent_report = parent_report
        self.year_id = self.parent_report.parent_project.year_id
        self.wea_hours = int(wea_hours) if wea_hours is not None else None
        self.wea_output = int(wea_output) if wea_output is not None else None
        self.inspectors_list = inspectors_list
        try: self.date = self.inspectors_list[0][2]
        except TypeError: self.date = None
        self.testrun_wind = float(testrun_wind) if testrun_wind is not None else None
        self.testrun_power = float(testrun_power) if testrun_power is not None else None
        self.timeline = timeline
        self.parent_report.parent_project.save()

    def __repr__(self):
        if self.has_happened:
            happened = f'happened on {self.get('date')}'
        else:
            happened = 'not yet happened'
        return repr(f'{self.kind} Inspection that has {happened}')

    def mark_happened(self):
        # convert data types when called
        self.testrun_power = gui_f.str2float(self.testrun_power)
        self.testrun_wind = gui_f.str2float(self.testrun_wind)
        
        self.has_happened = True
        self.insert_to_db()
        self.insert_to_db()
        # BUG: date in db changes to first date of first inspector only when saving the data a second time :o
        # caused by inspection.get_date() using the database entry, not the date of the inspectioupdate object

        self.parent_report.parent_project.save()

    def insert_to_db(self):
        '''inserts or updates database entry for this inspection'''
        # don't get values via .get method, because they come from the database
        # and thus won't be updated
        inspections_db = dbf.load_inspections()
        wea = self.parent_report.parent_wea
        inspections_db.loc[self.get_year_id(),
                            self.kind,
                            wea.get('oem'),
                            wea.id] = [f'{self.inspectors_list}',
                                       self.get_testrun_wind(),
                                       self.get_testrun_power(),
                                       self.get_wea_output(),
                                       self.get_wea_op_hours(),
                                       self.scope,
                                       self.timeline]
        dbf.save_db(inspections_db, 'inspections')

    def get_db_entry(self):
        inspections_db = dbf.load_inspections()
        year_id = self.get_year_id()
        wea = self.parent_report.parent_wea
        return inspections_db.loc[year_id, self.kind, wea.oem, wea.id]    

    def get_year(self):
        i = 0
        for character in self.year_id:
            if not character.isnumeric():
                break
            i += 1
        year_raw = self.year_id[:i]
        if len(year_raw) == 2:
            year = f'20{year_raw}'
        elif len(year_raw) == 4:
            year = year_raw
        else:
            raise ValueError((f'year substring ({year_raw}) of year_id '
                              f'{self.year_id} not interpretable as year'))
        return year
    
    def get(self, attr):
        if attr == 'date': return self.get_date()
        if attr == 'inspectors': return self.get_inspectors()
        if attr == 'testrun_wind': return self.get_testrun_wind()
        if attr == 'testrun_power': return self.get_testrun_power()
        if attr == 'wea_output': return self.get_wea_output()
        if attr == 'wea_op_hours': return self.get_wea_op_hours()
        if attr == 'kind': return self.kind.upper()
        if attr == 'timeline': return self.get_db_entry()['timeline']
        return self.__getattribute__(attr)

    def get_year_id(self):
        return self.year_id
    def get_date(self):
        insp = self.inspectors_list
        try: date_raw = insp[0][2]
        except TypeError: return None
        date_first = date_raw.split(',')[0]
        self.date = date_first.split(' ')[0]
        return self.date
    def get_inspectors(self):
        db = self.get_db_entry()
        try: inspectors_list = eval(db.inspectors)
        except TypeError: inspectors_list = db.inspectors
        return inspectors_list
    def get_testrun_wind(self):
        if self.testrun_wind == '' or self.testrun_wind is None:
            return None
        return float(self.testrun_wind)
    def get_testrun_power(self):
        if self.testrun_power == '' or self.testrun_power is None:
            return None
        return float(self.testrun_power)
    def get_wea_output(self):
        return self.wea_output if self.wea_output is None or self.wea_output=='' \
                    else int(self.wea_output)
    def get_wea_op_hours(self):
        return self.wea_hours if self.wea_hours is None or self.wea_hours=='' \
                    else int(self.wea_hours)
    
    def get_missing_data(self):
        missing = []
        insp = self.get_inspectors()

        if not insp:
            missing.append('Keine Prüfer angegeben')
            return missing

        if len(insp) == 1:
            missing.append(f'Nur 1 Person an Anlage angegeben')

        alldates = [ins[2] for ins in insp]
        allweathers = [ins[3] for ins in insp]

        if len(set(alldates)) == 1 and '' in set(alldates):
            missing.append('Kein Prüfdatum angegeben')
        if len(set(allweathers)) == 1 and '' in set(allweathers):
            missing.append('Keine Witterung angegeben')
        has_temp = False
        has_wind = True
        for weather in allweathers:
            if "m/s" in weather: has_wind = True
            if '°C' in weather: has_temp = True
        
        if not has_temp: missing.append('Keine Temperatur bei Prüfung angegeben')
        if not has_wind: missing.append('Keine Windgeschweindigkeit bei Prüfung angegeben')

        return missing
    
    def set_timeline_data(self, tl_dict: dict, start_month: str, end_month: str):
        '''set timeline data'''
        # checks
        if not isinstance(tl_dict, dict):
            raise ValueError(f'tl_dict mus be type dict but is {type(tl_dict)}:\n{tl_dict}')
        dbf.monthyear2datetime(start_month)
        dbf.monthyear2datetime(end_month)

        self.timeline = f'({tl_dict}, {start_month}, {end_month})'
        self.insert_to_db()


    

class Remark:
    '''single remark in report, class for inserting remark to document.'''
    def __init__(self,
                 address: str,
                 flag: str,
                 text: str,
                 parent_report,
                 image_names: Optional[list] = [],
                 timestamp: Optional[pd.Timestamp] = None,
                 ):
        '''
        address: unique identifier, form: chapter|section|...|title
        flag: V, P, I, 
            * (item in list without bullet),
            - (item in list with bullet), 
            S (regular sentence with paragraph at the end),
            0, 2, 3, 4 (colors gray, green, yellow, red)
        text: full text to put in remarks_db, text of the remark body in report
        image_paths: list with images associated with the remark, images must
                        sit in 1-Bericht'''
        self.address = address
        self.flag = flag.upper()
        self.text = text
        self.image_names = image_names
        self.parent_report = parent_report
        if not timestamp:
            self.timestamp = dbf.get_curr_time()
        else:
            self.timestamp = timestamp

    def __repr__(self):
        displaytext = self.text
        if len(displaytext) > 24:
            displaytext = f'{displaytext[:20]}...'
        return repr(f'{displaytext} in {self.address} with {self.flag} flag')

    def insert_to_document(self, doc: docx.Document=None, table: docx.table=None):
        '''format the remark according to flag and insert into the document'''
        if self.flag in 'PPPVEI': self._remark2doc(table)        
        elif self.flag in '-*S': self._sentence2doc(doc)
        elif self.flag in '0234': self.conclusionheader2doc(doc)
        elif self.flag == 'RAW': self._RAW2doc(doc)
        else: raise ValueError(f'remark flag unknown: {self.flag}. Please choose flag from dropdown.')

    def _remark2doc(self, table):
        imgwidth = 6.25 # cm
        row = table.add_row()
        gui_f.delete_paragraph(row.cells[0].paragraphs[0])
        paragraph = row.cells[0].add_paragraph(style='Listenabsatz')
        gui_f.simplynumber(paragraph)
        row.cells[1].text = self.flag
        gui_f.delete_paragraph(row.cells[2].paragraphs[0])
        for par in self.text.split(db_split_char):
            row.cells[2].add_paragraph(par)

        if not isinstance(self.image_names, list):
            if pd.isna(self.image_names):
                self.image_names = []

        for i, image in enumerate(self.image_names):
            if i%2 == 0:
                paragraph = row.cells[2].add_paragraph()
                sep = '' if i < 1 else '\n'
            else: sep = '      '
            paragraph.add_run(sep).add_picture(
                f'{self.parent_report.parent_wea.id}/1-Bericht/{image}',
                width=Cm(imgwidth if 'timeline' not in image else imgwidth*2))
        
        row.cells[0].width = Cm(1.6)
        row.cells[1].width = Cm(.75)
        row.cells[2].width = Cm(13.6)

    def _sentence2doc(self, doc):
        if self.flag == 'S': style = None
        elif self.flag == '-': style = 'Auflistung'
        elif self.flag == '*': style = 'Unsichtbare Auflistung'

        for par in self.text.split(db_split_char):
            doc.add_paragraph(par, style=style)

        if not isinstance(self.image_names, list):
            if pd.isna(self.image_names):
                self.image_names = []
        if not self.image_names: return
        print('Bild außerhalb von PEIV Bemerkung. Wird ignoriert.')

    def conclusionheader2doc(self, doc):
        # Create a paragraph for heading
        p = doc.add_paragraph()
        p.style = 'Heading 3'
        p.add_run().add_picture(f'{mainpath}/images/rectangle_{self.flag}.png',
                            width=Cm(.8), height=Cm(.35))
        p.add_run(f' {self.text}')

    def _RAW2doc(self, doc):
        # assume we are at the end of the doc in its current form
        # Create a namespace combining globals() and locals()
        namespace = {**globals(), **locals()}  # Merges global and local variables
        code_str = self.text.replace(db_split_char, '\n')
        exec(code_str, namespace)  # Execute in the isolated namespace



class Report:
    def __init__(self,
                 parent_wea,
                 parent_project,
                 inspection_kwargs: dict,
                 show_on_title: dict=None,
                 id: Optional[str] = None,
                 authors: Optional[list] = None,
                 ):
        '''authors: [(role, name, location, date, signature_image), (role, ...), ...]'''
        self.parent_wea = parent_wea
        self.parent_project = parent_project
        try:
            self.set_authors(authors)
        except TypeError:
            self.authors = None
        try:
            self.author = self.authors[0][1]
        except TypeError:
            self.author = config.default_author
        self.id = id
        self.inspection = Inspection(parent_report=self, **inspection_kwargs)
        self.set_chapters_from_checklist()
        self.init_chapter_renames()
        self.update_checklist()
        self.done_chapters = []
        self.parent_project.save()
        self.show_on_title = show_on_title

    def __repr__(self):
        if self.id is None:
            id_str = 'no ID'
        else:
            id_str = f'ID {self.id}'
        remarks = self.get_remarks(ordered=False)
        if remarks is None:
            remarks = []
        return repr(f'Report with {id_str} and {len(remarks)} remarks')
    
    def set_authors(self, authors):
        if isinstance(authors, list):
            self.authors = authors
            return
        elif isinstance(authors, str):
            self.authors = eval(authors)
            self.set_authors(self.authors)
        elif isinstance(authors, tuple):
            self.authors = list(authors)
            self.set_authors(self.authors)
        else:
            raise TypeError((f'Authors must be givan as [(role, name, location, '
                             'date, signature_image), (role, ...), ...], got '
                             f'{authors} instead.'))
        
    def get_authors(self):
        return self.authors
    
    def get_year(self):
        raw_date = self.get_latest_date_from_authors()
        raw_year =  raw_date[raw_date.rfind('.')+1:]

        try:
            if len(raw_year) == 2:
                return int(f'20{raw_year}')
            elif len(raw_year) == 4:
                return int(raw_year)
        except ValueError:
            pass
        # self.authors = None
        raise IndexError(
            f'Date of report must consist of 2 or 4 numbers (aktuell: {raw_year})')
    
    def get_latest_date_from_authors(self):
        available_dates = []
        for _, _, _, date, _ in self.authors:
            if not date:
                continue
            try:
                available_dates.append(pd.to_datetime(date, dayfirst=True))
            except:
                continue
        try:
            latest = pd.Series(available_dates).max().date()
        except AttributeError: # happens when none of available dates is date
            return 'DATEN AN DEN UNTERSCHRIFTEN KORRIGIEREN'
        return latest.strftime('%d.%m.%Y')
    
    def get_dates(self, raise_errors=True):
        '''raise_errors: bool, how to handle errors when parsing given dates'''
        available_dates = []
        for _, name, _, date, _ in self.authors:
            if not date and raise_errors:
                raise ValueError(f'no date for {name} given.')
            if not date: continue
            try:
                available_dates.append(pd.to_datetime(date, dayfirst=True))
            except Exception as e:
                if not raise_errors: continue
                raise e

    
    def create_folders(self):
        secs = ['0-Fertig', '1-Bericht', '2-Erwähnt', '3-Ignoriert', '4-Info',
                '5-Kennzeichnung', '6-Doku', '7-Typenschilder']
        for sec in secs:
            try:
                os.mkdir(f'{os.getcwd()}/{self.parent_wea.id}/{sec}')
            except FileExistsError:
                pass
    
    def add_remark(self, address, flag, text, pos_nr='',
                   timestamp: Optional[pd.Timestamp|None]=None,
                   image_names: Optional[list|None]=None):
        '''
        adds remark to report, saves it to database,
        remark_kwargs = [address, flag, text, image_names]
        '''
        remarks_db = dbf.load_remarks()
        if address[-1] == db_split_char:
            raise IndexError(f'address should not end with {db_split_char}')
        wea = self.parent_wea
        inspection = self.inspection
        oem = wea.oem
        id = str(wea.id)
        year_id = inspection.get_year_id()
        inspection_type = inspection.kind
        if timestamp is None:
            timestamp = dbf.get_curr_time()
        else:
            timestamp = pd.Timestamp(timestamp)
        if image_names is None:
            image_names=[]
        remarks_db.loc[oem,
                       id,
                       year_id,
                       inspection_type,
                       address,
                       timestamp] = [text,
                                     flag,
                                     image_names,
                                     self.author,
                                     pos_nr]
        dbf.save_db(remarks_db, 'remarks')

    def remove_remark(self, address, timestamp):
        '''removes remark from report and database
        address: str, full address chapter|section|...|title'''
        remarks_db = dbf.load_remarks()
        timestamp = pd.Timestamp(timestamp)
        wea = self.parent_wea
        year_id = self.inspection.get_year_id()
        index = (wea.oem, str(wea.id), year_id, self.inspection.kind,
                 address, timestamp)

        remarks_db.drop(index, inplace=True)
        dbf.save_db(remarks_db, 'remarks')

    def get_titles(self, remarks: Optional[pd.DataFrame]=None):
        if remarks is None: remarks = self.get_remarks(ordered=False)
        address_notitles = (remarks
                            .index
                            .get_level_values('address')
                            .str
                            .split(db_split_char)
                            .str[:-1]
                            .str
                            .join(db_split_char))
        titles = (remarks
                  .index
                  .get_level_values('address')
                  .str
                  .split(db_split_char)
                  .str[-1])
        return pd.Series(titles, index=address_notitles, name='titles')

    def get_remarks(self, index_startswith: Optional[str]='',
                    include_subchapters=True,
                    index_startswith_ends_with: Optional[str|None]='section',
                    ordered=True):
        '''returns all remarks whose address begins with index_startswith
        when using index_startswith and include_subchapters, index_startswith
        should not end with |
        index_startswith_ends_with: str section, title or None, default section.
            Only return entries whose address elements (section/title) are no
            longer than the section/title given in index_startswith'''
        if index_startswith_ends_with not in ['section', 'title', None]:
            raise ValueError('Argument index_startswith_ends_with must be section, '
                             'title or None.')

        try:
            park_rems = self.parent_project.get_all_remarks()
            report_remarks = dbf.filter_specific_report(park_rems,
                                                        self.parent_wea)
        except KeyError: return dbf.get_empty_remarks_df()
        if not index_startswith: return self.order_remarks(report_remarks) if ordered else report_remarks
        # handle configuration only remarks of certain chapter
        if index_startswith and not include_subchapters and\
             index_startswith_ends_with == 'section':
            chapter_remarks = dbf.get_remarks_of_chapter(report_remarks,
                                                         index_startswith)
            return self.order_remarks(chapter_remarks) if ordered else chapter_remarks

        correct_start = (report_remarks
                        .index
                        .get_level_values('address')
                        .str
                        .startswith(index_startswith))
        report_remarks = report_remarks.loc[correct_start]
        if not index_startswith_ends_with == 'title':
            return self.order_remarks(report_remarks) if ordered else report_remarks
        
        correct_length = (report_remarks
                          .index
                          .get_level_values('address')
                          .str
                          .len() == len(index_startswith))
        report_remarks = report_remarks.loc[correct_length]
        return self.order_remarks(report_remarks) if ordered else report_remarks
        

    def add_default_remarks(self, remarks=None):
        if not remarks: remarks = self.get_remarks(ordered=False)
        default_remarks = dbf.get_default_remarks(self.get_checklist())
        for index in default_remarks.index:
            db_remark = default_remarks.loc[[index]]
            _address = db_remark.index.get_level_values('chapter')[0]
            _title = db_remark.index.get_level_values('title')[0]
            address = f'{_address}|{_title}'
            flag = db_remark.default_state.iloc[0] # Series with 1 entry.
            text = db_remark.fulltext.iloc[0]
            # check if remark is already in the report.
            if address in remarks.index.unique('address'):
                texts = remarks.loc[address].fulltext.values
                if text in texts:
                    print(f'Remark \'{address}\' already in report. Won\'t be added again.')
                    continue
            self.add_remark(address=address, flag=flag, text=text)

    def mark_default_done_chapters_as_done(self):
        for chap in config.default_done_chapters:
            if chap in self.get_chapters() and chap not in self.get_done_chapters():
                self.done_chapters.append(chap)

    def order_remarks(self, remarks: pd.DataFrame):
        '''orders remarks according to chapter_order.txt
        also handle manually assigned positions'''
        ordered_addresses = dbf.dict2addresses(dbf.get_order('chapters'))
        ordered_chapters = (pd.Series(range(len(ordered_addresses)),
                                      index = ordered_addresses)
                                      .index
                                      .str
                                      .split(db_split_char)
                                      .str[:-1]
                                      .str
                                      .join(db_split_char)
                                      .to_list())
        present_addresses = dbf.ser2addresses(self.get_titles(remarks))

        # find addresses that are not in the chapter order dict and sort them in.
        # for now sorting means putting it at the chapter's start...
        strange_addresses = list(set(present_addresses) - set(ordered_addresses))

        for address in strange_addresses:
            chapter = address[:address.rfind(db_split_char)]
            i = ordered_chapters.index(chapter)
            ordered_addresses.insert(i, address)
        
        ordered_df = remarks.reset_index()
        ordered_df['address'] = pd.Categorical(ordered_df['address'],
                                               categories=ordered_addresses,
                                               ordered=True)
        ordered_df = ordered_df.sort_values('address')
        ordered_df = ordered_df.set_index(['address', 'create_time'])
        ordered_df = dbf.order_by_position(ordered_df)
        return ordered_df

    def get_inspection(self):
        return self.inspection
    
    def order_chapters(self):
        allchapters = list(set(chap for chap in self.chapters))
        ordered_list = []

        # sort chapters according to dbf.get_order('chapters') in lengthy_imports
        for chapter in dbf.get_order('chapters').keys():
            if chapter in allchapters:
                ordered_list.append(chapter)
                allchapters.remove(chapter)
        
        # put in remaining chapters at end of superchapter entries
        for remaining_chapter in allchapters:
            superchapter = remaining_chapter[:remaining_chapter.rfind(db_split_char)]
            pos = len(ordered_list)
            for i, chapter in enumerate(ordered_list):
                if chapter.startswith(superchapter):
                    pos = i+1
            ordered_list.insert(pos, remaining_chapter)
        if len(ordered_list) == len(self.chapters):
            self.chapters = ordered_list
            self.parent_project.save()
            return
        raise ValueError('something went wrong ordering the chapter list :(')
    
    def get(self, attr):
        if attr == 'chapter_renames':
            return self.chapter_renames
        if attr in self.__dict__.keys():
            return self.__dict__[attr]

    @gui_f.timeit
    def update_compressed_images(self, check_only=None):
        '''
        update images in 1-Bericht. Checks for changes in images in 0-Fertig und
        changes compressed images in 1-Bericht accordingly.
        check_only: list/None, list of image names that shall be checked, if None
        all image names in 0-Ferig are checked
        '''
        def get_img_diff() -> dict:
            '''get a dict with
            images that are in 0-Fertig but not in 1-Bericht (key add),
            images that are in 1-Bericht but not in 0-Fertig (key del),
            images that are in both but have different hashes (key change)
            
            handles special case of no file called hashes.txt in 1-Bericht
            update hashes.txt in 1-Bericht'''
            img_names_fertig = os.listdir(f'{os.getcwd()}/{self.parent_wea.id}/0-Fertig')
            img_names_compressed = os.listdir(f'{os.getcwd()}/{self.parent_wea.id}/1-Bericht')
            
            img_diff = {'add': [], 'del': [], 'change': []}
            img_hashes = {}
            hash_file = f'{os.getcwd()}/{self.parent_wea.id}/1-Bericht/hashes.txt'
            
            if os.path.exists(hash_file):
                with open(hash_file, 'r') as f:
                    for line in f:
                        name, hash = line.strip('()\n').split(', ')
                        img_hashes[name] = hash

            for img in img_names_fertig:
                if check_only is not None:
                    if img not in check_only: continue
                img_path = f'{os.getcwd()}/{self.parent_wea.id}/0-Fertig/{img}'
                img_hash = hl.md5(open(img_path, 'rb').read()).hexdigest()
                if img not in img_hashes:
                    img_diff['add'].append(img)
                elif (img_hashes[img] != img_hash) or (img not in img_names_compressed):
                    img_diff['change'].append(img)
                img_hashes[img] = img_hash

            for img in img_names_compressed:
                if check_only is not None:
                    if img not in check_only: continue
                if img == 'hashes.txt': continue
                if img not in img_hashes:
                    img_diff['del'].append(img)

            with open(hash_file, 'w') as f:
                for img, hash in img_hashes.items():
                    f.write(f'({img}, {hash})\n')

            return img_diff
        
        def add_image(img_name):
            img_path = f'{os.getcwd()}/{self.parent_wea.id}/0-Fertig/{img_name}'
            save_dir = f'{os.getcwd()}/{self.parent_wea.id}/1-Bericht/{img_name}'
            width = config.default_image_width_pxl if 'timeline' not in img_name \
                                else config.default_image_width_pxl*2
            gui_f.compress_image(img_path, save_to=save_dir, new_width_pxl=width)

        
        img_diff = get_img_diff()
        for img in img_diff['del']:
            os.remove(f'{os.getcwd()}/{self.parent_wea.id}/1-Bericht/{img}')
        
        for img in img_diff['add'] + img_diff['change']:
            add_image(img)

    def to_word(self):
        # CONVENTION: manually generated pagebreaks BEFORE chapters!
        def db2word(doc: docx.Document, db_slice: pd.DataFrame,
                    write_chapter_headings=False, pagebreak_before_chapter=True,
                    old_chap='', old_sec='', old_subsec='') -> docx.Document:
            '''wordifies slice of db. Optionally handle chapters.
            index mus be (address, create_time) - no more no less
            db_slice needs to be ordered!
            Adds chapters/sections/subsections to return string if
            write_chapters=True.
            old_chap/sec/subsec: headlines are written, once chap/sec/subsec in
            address differs from these
            '''

            if db_slice.empty: return

            def add_chapterheading(doc, chapter):
                if pagebreak_before_chapter: doc.add_page_break()
                doc.add_heading(chapter, 1)
            def add_sectionheading(doc, section):
                doc.add_heading(section, 2)
            def add_subsectionheading(doc, subsection, address):
                if address.startswith('Prüfbemerkungen|') and\
                     not config.show_subsections_in_remarks:
                    return
                doc.add_heading(subsection, 3)

            table_active = False
            old_address_no_title = ''

            renames = self.get('chapter_renames')
            
            if not isinstance(db_slice, pd.DataFrame) \
                    or isinstance(db_slice, pd.Series):
                return doc
            for address, create_time in db_slice.index:        # index is tuple
                address_no_title = address[:address.rfind(db_split_char)]
                if write_chapter_headings:
                    section = address[:address.rfind(db_split_char)]
                    if section in renames.keys():
                        newtitle = renames[section]
                        section = section[:section.rfind(db_split_char)+1]+newtitle

                    section_parts = section.split(db_split_char)
                    chap = section_parts[0]
                    try: sec = section_parts[1]
                    except IndexError: sec = ''
                    try: subsec = section_parts[2]
                    except IndexError: subsec = ''

                    if chap != old_chap:
                        add_chapterheading(doc, chap)
                        if sec: add_sectionheading(doc, sec)
                        if subsec: add_subsectionheading(doc, subsec, address)
                    elif sec and (sec != old_sec):
                        add_sectionheading(doc, sec)
                        if subsec: add_subsectionheading(doc, subsec, address)
                    elif subsec and (subsec != old_subsec):
                        add_subsectionheading(doc, subsec, address)
                    old_chap, old_sec, old_subsec = (chap, sec, subsec)


                text, flag, images, _, _ = db_slice.loc[address, create_time].to_list()
                needs_table = True if flag in 'PPPEIV' else False
                new_sec = True if address_no_title != old_address_no_title else False
                if isinstance(images, str):
                    images = eval(images)

                rem = Remark(address, flag, text, self, images,
                             pd.Timestamp(create_time))
                if needs_table:
                    if not table_active or new_sec:
                        remtable = doc.add_table(rows=0, cols=3, style='Bemerkungstabelle')
                        remtable.autofit = False
                    rem.insert_to_document(table=remtable)
                    table_active = True
                else:
                    rem.insert_to_document(doc=doc)
                    table_active = False
                old_address_no_title = address_no_title
            return doc
        
        def titlepage2doc(doc):
            vspace_long = 1.5
            vspace_short = .5

            wea = self.parent_wea
            wea_db = wea.get_db_entry()


            t = doc.add_table(cols=2, rows=0, style='Bemerkungstabelle')
            titleinfo = {
                'Aufgabenstellung': inspection_type_translator[self.inspection.get('kind')],
                'Prüfumfang': self.inspection.get('scope'),
                'vspace0': vspace_long,               # height of blank vertical space in cm
                'WEA-Typ': f'{wea.get('oem')} {wea_db.model} '
                           f'{wea_db.submodel\
                               if not pd.isna(wea_db.submodel) else ''}',
                'Seriennummer': wea.get('id')
                }
            # handle potentially missing farm number
            farm_number = wea_db.farm_number
            if pd.isna(farm_number) or farm_number == '' or farm_number == None:
                farm_number = None
            if (farm:=wea_db.windfarm) and farm_number:
                titleinfo['Windpark und Standortnr.'] = f'{farm}, {farm_number}'
            elif farm:
                titleinfo['Windpark'] = farm
            titleinfo['Standort der WEA'] = wea_db.location

            titleinfo['vspace1'] = vspace_long
            titleinfo['Betreiber'] = wea_db.owner

            if contact:=self.parent_project.get('contact'):
                titleinfo['vspace2'] = vspace_short
                titleinfo['Auftraggeber'] = contact
            titleinfo['vspace3'] = vspace_short
            titleinfo['Auftragsdatum'] = self.parent_project.get('order_date')
            titleinfo['vspace4'] = 2
            titleinfo['Auftragnehmer'] = self.parent_project.get('contractor')

            if subcontractor:=self.parent_project.get('subcontractor'):
                titleinfo['vspace5'] = vspace_short
                titleinfo['Unterauftragnehmer'] = subcontractor
            titleinfo['vspace6'] = vspace_long
            titleinfo['Verantwortl. Ingenieur'] = self.parent_project.get('engineer')
            titleinfo['vspace7'] = vspace_long
            titleinfo['Datum'] = self.get_latest_date_from_authors()

            for what, value in titleinfo.items():
                if what.startswith('vspace'):
                    row = t.add_row()
                    row.height = Cm(value)
                    continue
                row = t.add_row().cells
                row[0].text = f'{what}:'
                # new paragraph for each newline. so that blocksatz is not spanning entire page despite short text
                c = row[1]
                p = c.paragraphs[0]
                gui_f.delete_paragraph(p)
                for text_fragment in value.split(db_split_char):
                    # filter out mailaddress (works only if format is ' mail@provider' and then stop. note the blank before mail)
                    if mail:=('@' in text_fragment):
                        cutoff = text_fragment.rfind(' ')+1
                        text_fragment_no_mail = text_fragment[:cutoff]
                        mailaddress = text_fragment[cutoff:]
                        text_fragment = text_fragment_no_mail

                    p = c.add_paragraph(text_fragment)
                    if mail: gui_f.add_hyperlink(p, mailaddress, f'mailto:{mailaddress}')
                    

            gui_f.set_table_column_widths(t, [.35, .65], relative=True, doc=doc)


        def TOC2doc(doc):
            doc.add_page_break()
            doc.add_paragraph('Inhaltsverzeichnis', style='Überschrift Inhalt')
            p = doc.add_paragraph()
            r = p.add_run()
            gui_f.add_TOC(r)

        def generaldata2doc(doc: docx.Document):
            doc.add_page_break()            

            doc.add_heading('Allgemeine Daten', 1)
            doc.add_heading('Daten zur Prüfung', 2)

            # insert inspectors/dates/weather...
            t = doc.add_table(rows=1, cols=4, style='Minimale Tabelle')
            header = t.rows[0].cells
            for i, text in enumerate(['Aufgabe', 'Name', 'Datum', 'Witterung']):
                header[i].text = text

            inspectors = self.inspection.get_inspectors()
            for single_inspector_data in inspectors:
                row = t.add_row().cells
                for i, data in enumerate(single_inspector_data):
                    row[i].text = data
            gui_f.set_table_column_widths(t, [.2, .35, .2, .2], relative=True, doc=doc)
            gui_f.leftalign_table(t)

            # create remaining tables/sentences...
            rems = self.get_remarks('Allgemeine Daten')
            db2word(doc, rems,
                    write_chapter_headings=True, old_chap='Allgemeine Daten')

        def turbinedata2doc(doc):
            doc.add_page_break()            

            inspection_db = self.inspection.get_db_entry()
            turbine_db = self.parent_wea.get_db_entry()

            
            # turbine data table
            mean_power = inspection_db.wea_output/inspection_db.wea_hours
            insertions = [('Nennleistung',          int(float(turbine_db.rated_power)),                 'kW'),
                          ('Nabenhöhe',             int(float(turbine_db.hub_height)),                  'm'),
                          ('Rotordurchmesser',      int(float(turbine_db.rotor_diam)),                  'm'),
                          ('Inbetriebnahmedatum',   turbine_db.startup_date                                 ),
                          ('Betriebsstunden',       gui_f.format_large_ints(inspection_db.wea_hours),   'h'),
                          ('Ertrag',                gui_f.format_large_ints(inspection_db.wea_output),  'kWh'),
                          ('Durchschnittsleistung', int(float(mean_power)),                             'kW'),
                          ]

            doc.add_heading('Anlagendaten', 1)
            doc.add_heading('Anlagen- und Betriebsdaten', 2)
            t = doc.add_table(cols=3, rows=0, style='Bemerkungstabelle')
            for _insertion in insertions:
                row = t.add_row().cells
                for i, value in enumerate(_insertion):
                    row[i].text = str(value)
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            gui_f.set_table_column_widths(t, [7, 3, 2])

            # turbine parts table
            doc.add_heading('Daten der Hauptkomponenten', 2)
            parts = self.parent_wea.get_parts(sorted=True)
            t = doc.add_table(rows=0, cols=3, style='Minimale Tabelle')
            row = t.add_row().cells
            for i, header in enumerate(['Bauteil', 'Typ', 'Nr. / Anzahl']):
                row[i].text = header
                
            for partname in parts.keys():
                row = t.add_row().cells
                oem, model, sn = parts[partname]
                oem = f'{oem}, '
                if oem.lower().strip() == 'nan,':
                    oem = ''
                    model = ''
                if model.lower().strip() == 'nan':
                    model = ''
                    # remove trailing comma after oem if no model is given
                    if oem:
                        oem = oem[:-2]
                if sn == 'nan':
                    sn = ''
                
                row[0].text = partname
                row[1].text = f'{oem}{model}'
                row[2].text = f'{sn}'

            gui_f.leftalign_table(t)
            gui_f.set_table_column_widths(t, [.25, .4, .35], relative=True, doc=doc)
            # insert extras
            extras = self.get_remarks('Anlagendaten')
            db2word(doc, extras, write_chapter_headings=True, old_chap='Anlagendaten')


        def inspectionbase2doc(doc):
            db2word(doc, self.get_remarks('Prüfungsgrundlage'), write_chapter_headings=True)

        def remarks2doc(doc):
            db2word(doc, self.get_remarks('Prüfbemerkungen'), write_chapter_headings=True)

        def conclusion2doc(doc):
            # first part
            db2word(doc, self.get_remarks('Prüfergebnis|Fazit'),
                    write_chapter_headings=True)
            doc.add_page_break()
            db2word(doc, self.get_remarks('Prüfergebnis|Auflagen bzw. weiteres Vorgehen'),
                    write_chapter_headings=True, old_chap='Prüfergebnis')
            
            # add signature table
            authors = copy.deepcopy(self.authors)
            doc.add_paragraph() # for more distance
            t = doc.add_table(rows=0, cols=2, style='Bemerkungstabelle')

            for i, (role, name, place, date, signature) in enumerate(authors):
                col = 1
                if i % 2 == 0:
                    col = 0
                    for j in range(4):
                        t.add_row()
                line = int(i/2)
                role_cell = t.rows[line*4].cells[col]
                signature_cell = t.rows[line*4+1].cells[col]
                name_cell = t.rows[line*4+2].cells[col]
                placedate_cell = t.rows[line*4+3].cells[col]

                role_cell.text = role
                p = signature_cell.paragraphs[0]
                p.add_run().add_picture(f'{mainpath}/databases/report/{signature}', width=Cm(5.5))
                signature_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                name_cell.text = name
                placedate_cell.text = f'{place}, {date}'
                placedate_cell.add_paragraph(' ')
            

        self.update_compressed_images()
        
        rep_id = self.get('id')
        ingbuero_str = config.footnote_companyname if len(rep_id) < 40\
                    else config.footnote_companyname_abbrev

        doc = docx.Document(f'{mainpath}/8p2PBVorlage.docx')
        titlepage2doc(doc)
        TOC2doc(doc)
        gui_f.add_custom_footer(doc, ingbuero_str, rep_id)
        generaldata2doc(doc)
        turbinedata2doc(doc)
        inspectionbase2doc(doc)
        remarks2doc(doc)
        conclusion2doc(doc)

        doc_target = f'{os.getcwd()}/{self.parent_wea.get('id')}/{self.get('id')}.docx'
        doc.save(doc_target)
        

        
    def get_chapters(self, remarks=None):
        self.update_chapters()
        if remarks is None: remarks = self.get_remarks(ordered=False)
        rem_chaps = dbf.get_chapters_from_remarks(remarks)
        for chap in rem_chaps:
            if chap not in self.chapters:
                self.chapters.append(chap)
        self.order_chapters()
        return self.chapters
        
    def get_renamed_chapters(self, remarks=None):
        newnames = self.get_chapters(remarks)
        
        try: renames = self.get('chapter_renames')
        except AttributeError: renames = {}

        if not renames: return newnames

        for address in renames.keys():
            if address not in newnames:
                raise ValueError(f'address {address} is renamed, but not in the report\'s chapters.')
            newname = renames[address]
            newaddress = address[:address.rfind(db_split_char)+1] + newname
            newnames = [newaddress if oldaddress == address else oldaddress for oldaddress in newnames]
        return newnames

    def rename_chapter(self, old_address, new_title):
        og_title = old_address[old_address.rfind(db_split_char)+1:]
        if og_title == new_title:
            try: del self.chapter_renames[old_address]
            except KeyError: pass
            return
        self.chapter_renames[old_address] = new_title
        self.parent_project.save()
        
    def add_chapter(self, address):
        self.chapters.append(address)
        self.order_chapters()
        self.update_checklist()

    def update_chapters(self):
        self.update_checklist()
        chaps = dbf.get_chapters_from_checklist(self.checklist)
        chaps.extend(dbf.get_chapters_from_remarks(self.get_remarks(ordered=False)))
        self.chapters = list(set(chaps))
        self.order_chapters()
        self.parent_project.save()

    def get_timeline_data(self):
        '''get timeline dict, start month, end month if it is present in the
        db entry for this inspection.'''
        tl_data_raw = self.inspection.get('timeline')
        if pd.isna(tl_data_raw) or str(tl_data_raw).lower() in ['', 'nan']:
            return None
        tl_dict, start_month, end_month = eval(tl_data_raw)
        return tl_dict, start_month, end_month

    def update_checklist(self):
        '''needs to be called when a new remark is inserted into the complete_checklist,
        to reflect that change in the current report checklist'''
        self.checklist = dbf.get_report_checklist(self)
        self.parent_project.save()

    def get_checklist(self):
        self.update_checklist()
        return self.checklist

    def set_chapters_from_checklist(self):
        self.chapters = dbf.get_chapters_from_checklist(self.get_checklist())

    def init_chapter_renames(self):
        self.chapter_renames = copy.deepcopy(self.parent_project.get('chapter_renames'))
    
    def mark_chapter_done(self, chapter):
        if chapter not in self.chapters:
            raise ValueError(f"{chapter} not in the report's chapters")
        if chapter in self.get_done_chapters():
            return
        self.done_chapters.append(chapter)
        self.parent_project.save()

    def mark_chapter_undone(self, chapter):
        if chapter not in self.chapters:
            raise ValueError(f"{chapter} not in the report's chapters")
        if chapter not in self.get_done_chapters():
            return
        self.done_chapters.remove(chapter)
        self.parent_project.save()

    def get_chapter_done(self, chapter):
        if chapter in self.get_done_chapters():
            return True
        return False
    
    def get_first_undone_chapter(self):
        for chapter in self.chapters:
            if chapter not in self.get_done_chapters():
                return chapter
        return self.chapters[-1]
    
    def get_done_chapters(self):
        return self.done_chapters
    
    def get_missing_images(self, where='report', remarks=None) -> list:
        '''where: str, "report" or "folder". determines if the returned images
        are the images in 0-Fertig but missing in the report ("report") or if the
        images are present in report but missing from 0-Fertig
        '''
        if where not in ['report', 'folder']:
            raise ValueError(f'where must be "report" or "folder", but is {where}')
        images_in_remarks = self.get_image_list(remarks=remarks)
        
        images_in_folder = os.listdir(f'{os.getcwd()}/{self.parent_wea.id}/0-Fertig')
        if where == 'report':
            missing_imgs = [image for image in images_in_folder \
                                    if image not in images_in_remarks]
        elif where == 'folder':
            missing_imgs = [image for image in images_in_remarks \
                                    if image not in images_in_folder]
        return missing_imgs
    
    def get_multiply_used_imgs(self, remarks=None):
        used_imgs = pd.Series(self.get_image_list(remarks=remarks))
        img_counter = used_imgs.groupby(used_imgs).count()
        multiply_used = img_counter[img_counter > 1].index.to_list()
        return multiply_used

    def get_image_list(self, remarks=None):
        '''get all images (names, not paths) that are in the remarks for this 
        report as a list of str'''
        if remarks is None: remarks = self.get_remarks(ordered=False)
        if remarks is None: return []
        images_in_remarks = remarks.image_names.values
        images_in_remarks = [eval(image_list) for image_list in images_in_remarks]
        images_in_remarks = [image for image_list in images_in_remarks\
                                 if len(image_list)>0 for image in image_list]
        return images_in_remarks

    def get_missing_chapters(self, remarks=None) -> list:
        allchaps = self.get_chapters(remarks=remarks)
        donechaps = self.get('done_chapters')
        return [chap for chap in allchaps if chap not in donechaps]
    
    def get_missing_refs(self, remarks=None) -> list:
        missing_refs = []
        if remarks is None: remarks = self.get_remarks(ordered=False)
        if remarks is None: return missing_refs
        index_refs = [f'{ind[-2]}|{ind[-1]}' for ind in remarks.index]
        for ind in remarks.index:
            rem = remarks.loc[ind]
            text = rem.fulltext
            ref_begin = text.find('\\ref{')
            if ref_begin == -1:
                continue
            ref_begin += 5
            ref_end = text[ref_begin:].find('}')
            ref_str = text[ref_begin:][:ref_end]
            if ref_str in index_refs:
                continue
            missing_refs.append((f'Bemerkung wird referenziert, ohne zu existieren: '
                                 f'ref auf {ref_str} in Bemerkung {ind}'))
        return missing_refs
    
    def get_missing_attributes(self, clearname=False):
        missing = []
        for clearname in report_properties:
            kw = report_properties[clearname]
            attr = self.__getattribute__(kw)
            if attr: continue

            missing.append(clearname if clearname else kw)
        return missing
    
    def get_todo_count(self, remarks=None):
        if remarks is None: remarks = self.get_remarks(ordered=False)
        if remarks is None: return 0
        return remarks.fulltext.str.count(r'\?\?\?').sum()

    def completely_signed(self):
        try:
            self.get_dates()
        except:
            return False
        for author in self.authors:
            if '' in author:
                return False
            for val in author:
                if not val:
                    return False
                if pd.isna(val):
                    return False
        return True




class Turbine:
    '''class for single turbine, contains all information about the wea,
    generates its report'''
    def __init__(self, 
                 oem: str,
                 id: str,
                 is_setup = False,
                 report: Optional[Report] = None,
                 report_kwargs: Optional[dict] = None,
                 **turbine_kwargs
                 ):
        '''turbine_kwargs: model: str, rated_power: int, hub_height: float, 
        rotor_diam: float, tower_type: str, windfarm: str, location: str, owner: str,
        operator: str, startup_date: str, farm_number: int, coordinates: tuple,
        note: str
        '''

        #general info
        self.oem = oem.strip()
        self.id = id

        db_data = self.copy_db_data(return_data=True)

        # fill the rest...
        for attribute in turbine_kwargs.keys():
            if attribute not in db_data.index:
                print(f'{attribute} is not part of turbines.csv')
                continue
            value = self.__getattribute__(attribute)
            if value is None or value == '':
                self.__setattr__(attribute, turbine_kwargs[attribute])

        self.correct_data_types()

        # turbine specific info
        if report is not None:
            self.report = report
        else:
            self.report = Report(parent_wea=self, **report_kwargs)

        self.is_setup = is_setup
        self.report.create_folders()
        self.report.parent_project.save()


    def __repr__(self):
        id_str = f'ID {self.id}' if self.id is not None else 'no ID'
        try:
            model_str = f' {self.model}' if self.model is not None else ''
        except AttributeError:
            model_str = ''
        return repr(f'{self.oem}{model_str} with {id_str}')
    
    def setup(self):
        '''sets up the turbine, inserts it to db, inserts inspection to db'''
        if self.get('oem') in [None, '', 'non_setup']:
            raise ValueError(f'{self} has no valid oem')
        self.report.update_checklist()
        self.correct_data_types() 
        self.insert_to_db()
        self.report.inspection.mark_happened()  # setup happens after inspection, put in db only when oem is clear (after setup)
        self.report.parent_project.save()
        if self.is_setup: return

        self.report.add_default_remarks()
        self.report.mark_default_done_chapters_as_done()
        self.is_setup = True
                

    def get_report(self):
        return self.report
    
    def get_model(self):
        return f'{self.oem} {self.model}'
    
    def get(self, attr):
        db_entry = self.get_db_entry()
        if db_entry.isna().all(): return self.__getattribute__(attr)
        if attr == 'oem': return db_entry.name[0]
        if attr == 'id': return db_entry.name[1]
        try: val = db_entry[attr]
        except:
            IPS()
        if pd.isna(val): val = None
        return val

    def get_parts(self, **kwargs) -> dict:
        return dbf.get_turbine_parts(self.get('oem'),
                                     self.get('id'), **kwargs)
    
    def copy_db_data(self, return_data=False):
        db_data = self.get_db_entry()
        for attribute in db_data.index:
            value = db_data[attribute]
            if not pd.isna(value) and not value == '':
                self.__setattr__(attribute, value)
            else:
                self.__setattr__(attribute, None)
        if return_data:
            return db_data
    
    def insert_to_db(self):
        '''inserts or updates database entry for this turbine'''
        turbines_db = dbf.load_turbines()
        self.correct_data_types()
        turbine_data = [self.model,
                        self.submodel,
                        self.windfarm,
                        self.location,
                        self.farm_number,
                        self.rated_power,
                        self.hub_height,
                        self.rotor_diam,
                        self.tower_type,
                        self.startup_date,
                        str(self.coordinates),
                        self.owner,
                        self.operator,
                        self.note,
                        ]
        turbines_db.loc[(self.oem, self.id), :] = turbine_data
        dbf.save_db(turbines_db, 'turbines')

    def get_db_entry(self):
        db = dbf.load_turbines()
        try:
            db_entry = db.loc[self.oem, str(self.id)]
        except KeyError:
            return pd.Series(index=pd.Index(db.columns))
        if type(db_entry) == pd.DataFrame:
            if len(db_entry) > 1:
                raise KeyError(f'WEA {self} has multiple entries in turbines_db. '
                               'Please clean up.')
            return db_entry.iloc[0]
        elif type(db_entry) == pd.Series:
            return db_entry
        
    def str2coordinates(self, coordinates_str):
        if coordinates_str is None:
            return None
        if coordinates_str == '':
            return ''
        if type(coordinates_str) is tuple:
            return coordinates_str
        
        coordinates_str = eval(coordinates_str)

        if type(coordinates_str) is tuple:
            return coordinates_str
        
        coords_list = gui_f.str2list(coordinates_str, split_char=', ')
        return tuple(gui_f.str2float(coord) for coord in coords_list)
    
    def correct_data_types(self):
        for attribute in ['model', 'submodel', 'tower_type', 'windfarm', 'location',
                         'owner', 'operator', 'startup_date', 'note']:
            value = self.__getattribute__(attribute)
            if value is None or value=='':
                if attribute == 'note':
                    continue
                print(f'{self.id} Warning: no value for {attribute}')
                continue
            self.__setattr__(attribute, value.strip())
        for attribute in ['rated_power', 'farm_number']:
            value = self.__getattribute__(attribute)
            if value is None or value=='':
                continue
            if isinstance(value, str) and '.' in value:
                value = value[:value.find('.')]
            self.__setattr__(attribute, int(value))
        for attribute in ['hub_height', 'rotor_diam']:
            value = self.__getattribute__(attribute)
            self.__setattr__(attribute, gui_f.str2float(value))
        self.coordinates = self.str2coordinates(self.coordinates)

    def get_missing_attributes(self, clearname=False):
        
        missing = []
        for clearname in turbine_properties:
            if clearname == 'Kurznotiz' or clearname == 'Untermodell': continue
            kw = turbine_properties[clearname]
            try:
                attr = self.__getattribute__(kw)
            except AttributeError:
                attr = None
            if attr: continue
            missing.append(clearname if clearname else kw)
        return missing

class Windfarm:
    def __init__(self,
                 name: str='',
                 location: str='',
                 weas: Optional[dict]={}):        
        self.name = name
        self.location = location
        self.weas = weas

    def __repr__(self):
        return repr(f'Windfarm {self.name} at {self.location} '
                    f'with {len(self.weas)} turbines')

    def get_wea_ids(self) -> list:
        return list(self.weas.keys())

    def get_setup_wea_ids(self) -> list:
        setup_wea_ids = []
        for wea_id in self.weas.keys():
            wea = self.weas[wea_id]
            if wea.is_setup:
                setup_wea_ids.append(wea_id)
        return setup_wea_ids 

    def get_wea(self, wea_id) -> Turbine:
        try:
            return self.weas[wea_id]
        except KeyError:
            return None

    def add_wea(self, wea: Turbine):
        self.weas[wea.id] = wea

    def remove_wea(self, wea_id):
        del self.weas[wea_id]

    def get_first_non_setup_wea(self) -> Turbine:
        for wea in self.weas.values():
            if not wea.is_setup:
                return wea
        return None
    
    def get_setup_weas(self) -> list:
        setup_weas = []
        for wea in list(self.weas.values()):
            if wea.is_setup:
                setup_weas.append(wea)
        return setup_weas      

    def count_completed_inspections(self) -> int:
        counter = 0
        for wea in list(self.weas.values()):
            if wea.report.inspection.has_happened:
                counter += 1
        return counter
    
    def get_weas(self) -> list:
        return list(self.weas.values())

class Project:
    def __init__(self,
                 name: str,
                 year_id: str,
                 windfarm: Windfarm=Windfarm(),
                 order_date: str='',
                 contact: str='',
                 contractor: str='',
                 engineer: str='',
                 subcontractor: Optional[str]=None,
                 ) -> None:
        self.name = name
        self.year_id = year_id
        self.windfarm = windfarm
        self.order_date = order_date
        self.contact = contact
        self.contractor = contractor
        self.engineer = engineer
        self.subcontractor = subcontractor
        self.active_page = ('', '') # [wea_id, chapter]
        self.chapter_renames = {}


    def __repr__(self):
        if self.windfarm is not None:
            return repr(
                f'Projekt {self.get('name')} ({self.year_id}) '
                f'with {len(self.windfarm.weas)} Turbines in {self.windfarm.location} '
                f'({self.windfarm.count_completed_inspections()} inspected)')
        return repr(f'Projekt {self.get('name')} ({self.year_id})')

    def get(self, attr):
        if attr == 'name':
            return self.name # replace with code to infer name
        return self.__getattribute__(attr)
    
    def get_inspection_kind(self) -> str:
        wea = self.windfarm.get_first_non_setup_wea()
        if wea is None: wea = self.windfarm.get_setup_weas()[0]
        return wea.report.inspection.kind

        
    def set_active_wea(self, wea_id):
        old_chapter=self.get('active_page')[1]
        self.active_page = (wea_id, old_chapter)
        self.save()
    def set_active_chapter(self, chapter):
        old_wea=self.get('active_page')[0]
        self.active_page = (old_wea, chapter)
        self.save()

    def save(self):
        filename = 'project.pickle'
        with open(filename, 'wb') as f:
            pickle.dump(self, f)

    def create_notes(self):
        weas = self.windfarm.get_wea_ids()
        with open(f'notizen {self.name}.txt', 'a', encoding='utf-8') as f:
            f.write(f'{f':{3*'\n'}'.join(weas)}:{10*'\n'}')

    def enter_console(self):
        IPS()

    def create_overview(self, style: str='Repotex'):
        '''style: str, DataDocx or Inspect'''
        if style not in ['DataDocx', 'Inspect']:
            AttributeError(f'style must be DataDocx or Inspect, not {style}')
        if style == 'DataDocx':
            dbf.excel_overview(self)
        else: dbf.excel_overview_8p2_Inspect(self)
        print('successfully created an overview.')

    def get_missing_attributes(self, clearname=False):
        missing = []
        for clearname in project_properties:
            if clearname == 'Unterauftragnehmer': continue
            kw = project_properties[clearname]
            attr = self.__getattribute__(kw)
            if attr: continue
            missing.append(clearname if clearname else kw)
        return missing
    
    def rename_chapters_in_all_reports(self, old_address, new_title):
        self.chapter_renames[old_address] = new_title
        for wea in self.windfarm.get_weas():
            rep = wea.report
            rep.rename_chapter(old_address, new_title)
        self.save()

    def get_all_remarks(self, allowed_flags: Optional[list]=None,
                        drop_cols: Optional[list]=None) -> pd.DataFrame:
        return dbf.get_project_remarks(self, allowed_flags, drop_cols)

import os
import stat
import pickle
from lengthy_imports import *

import re
import subprocess

from PIL import Image, ImageTk, ImageOps

import numpy as np
import pandas as pd
import time

import docx
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ipydex import IPS, activate_ips_on_exception

activate_ips_on_exception()

def get_subfolder_names(projpath):
    subfolders = [x for x in os.walk(projpath)][0][1]
    return subfolders

def get_foldername(projpath):
    foldername = os.path.basename(os.path.normpath(projpath))
    return foldername

def get_folders_in_directory():
    list_subfolders = [f.name for f in os.scandir(os.getcwd()) if f.is_dir()]
    return list_subfolders

def get_turbtpye(wea):
    return f'{wea.oem} {wea.model}'

def get_attribute_from_clearname(object,
                                 clearname: str,
                                 translation: dict,
                                 ):
    '''returns an attrbute value of an object after translating the attribute
    name from a clearname to an attrbute name using a translation dict.
    If the object is a dict, it looks up clearname and argname in the dict and
    return whichever has a value'''
    attribute_value = None
    if isinstance(object, dict):
        if val := object.get(clearname, None): return val
        if val := object.get(translation[clearname], None): return val
        print(f'{object} has neither attribute {clearname} nor {translation[clearname]}, so I return None.')
    elif object and clearname and translation: # check if they are != None / 0
        argname = translation[clearname]
        try:
            attribute_value = object.__getattribute__(argname)
        except AttributeError as e:
            print(f'{object} has no attribute {argname}, so I return None.')
    return attribute_value

def change_attributes_from_dict(object_to_change, attr_dict):
    for argname in attr_dict.keys():
        value = attr_dict[argname]
        # test if object has attribute
        object_to_change.__getattribute__(argname)
        object_to_change.__setattr__(argname, value)
        
def load_project():
    try:
        with open('project.pickle', 'rb') as f:
            project = pickle.load(f)
    except FileNotFoundError:
        project=None
    return project

def argname2clearname(argname: str, translation: dict):
    for clearname, arg in translation.items():
        if arg == argname: return clearname
    raise AttributeError(f'argname {argname} not found in translation dict.')

def address_as_list(address):
    return address.split(db_split_char)
        
def str2list(convertible: str, truncate=' ', split_char=', '):
    if convertible is None or convertible == '':
        return None
    
    raw_list = convertible.split(split_char)
    if truncate is None:
        return raw_list
    return [element.strip() for element in raw_list]

def str2float(convertible: str):
    if convertible is None or convertible == '':
        return None
    if isinstance(convertible, str):
        convertible = convertible.replace(',', '.').strip()
        return float(convertible)
    elif isinstance(convertible, (int, float)):
        return float(convertible)
    else: raise AttributeError(f"can't convert to float, argument must be str, "
                               f"int or float, but is {type(convertible)} "
                               f"({convertible})")
    
def get_subdir_image_files(superdir):
    imagefiles = []
    for path, _, files in os.walk(superdir):
        for file in files:
            if file.endswith(('.jpg', 'jpeg', 'png')):
                imagefiles.append(os.path.join(path, file))
    return imagefiles

def coordinates_from_images(wea_id):
    def DMS2coords(degree, minutes, seconds, direction):
        coord = degree + minutes / 60 + seconds / 3600
        if direction in ['W', 'S']: coord *= -1
        return np.round(float(coord), 6)

    images = get_subdir_image_files(f'{os.getcwd()}/{wea_id}')
    x = []
    y = []
    for image in images:
        folder = image.split('\\' if sys.platform == 'win32' else '/')[-2]
        if folder in ['0-Fertig', '1-Bericht', '6-Doku']: continue
        with Image.open(image) as img:
            try:
                dms = img._getexif()[34853]
            except: continue
        try:
            x.append(DMS2coords(*dms[2], dms[1]))
            y.append(DMS2coords(*dms[4], dms[3]))
        except: continue
    if len(x) < 20 or len(y) < 20:
        raise AttributeError(f'Not enough images have GPS data.')
    x = np.median(np.array(x))
    y = np.median(np.array(y))
    return f'{x}, {y}'

def get_year(date: str):
    year = int(date[date.rfind('.')+1:])
    if len(str(year)) == 4 and str(year)[0] == '2':
        return year
    if len(str(year)) == 2:
        return int(f'20{year}')
    raise ValueError(f"cannot convert str to year-int with 4 digits and starting with '20': {year}")

def delete_children(widget, leave_out=None):
    '''delete a widgets children. leave out can by a type of widget or specific
    widget'''
    for w in widget.winfo_children():
        if leave_out:
            try:
                if isinstance(w, leave_out): continue
            except TypeError: pass
            if w == leave_out: continue
        w.destroy()

def grid_forget_children(widget):
    for w in widget.winfo_children():
        w.grid_forget()

def get_supermodel_from_model(model: str, cutoff_start: int=3):
    '''cuts off anything non-alphynumerical in a str after character at
    cutoff_start'''

    model = model.strip()
    cutoff_index = re.search(r'\W+', model[cutoff_start:])
    if cutoff_index is not None:
        cutoff_index = cutoff_index.start() + cutoff_start
        return model[:cutoff_index]
    return model

def get_strlist_union(*args, split_char=db_split_char):
    '''return unique items of the lists in args'''
    alllist = []
    for l in args:
        l = l.split(split_char)
        for i in l:
            try:
                i = i.strip()
            except:
                IPS()
            if i and i not in alllist:
                alllist.append(i)
    return alllist


def list2str(l):
    ret = ''
    for i in l:
        ret += f'{i}, '
    return ret


def get_tkinter_pic(image_path, width_pxl: int=160,
                    alias_fct=Image.LANCZOS, **PhotoImage_kw) -> ImageTk:
    try: pic = Image.open(image_path.replace('0-Fertig', '1-Bericht'))
    except FileNotFoundError: pic = Image.open(image_path)
    if width_pxl == 'original': resized = pic
    else:
        og_width, og_height = pic.size
        new_size = (width_pxl, int(width_pxl/og_width*og_height))
        resized = pic.resize(new_size, alias_fct)
    tk_pic = ImageTk.PhotoImage(resized, **PhotoImage_kw)
    return tk_pic

def open_image(image_path):
    open_file(image_path)

def rm_children(top):
    for root, dirs, files in os.walk(top, topdown=False):
        for name in files:
            filename = os.path.join(root, name)
            os.chmod(filename, stat.S_IWUSR)
            os.remove(filename)
        for name in dirs:
            os.rmdir(os.path.join(root, name))

def format_large_ints(int_to_format) -> str:
    '''separate number groups of 3 with a "."'''
    int_str = str(int(float(int_to_format))).strip()
    l = len(int_str)
    number_of_points = int((l-1)/3)
    used_points = 0
    pos = -3
    while used_points < number_of_points:
        int_str = int_str[:pos] + '.' + int_str[pos:]
        pos -= 4 # because point also as a position
        used_points += 1

    return int_str


def dbtext2displaytext(dbtext):
    text = dbtext.strip().replace('\r\n', '\n')
    text = re.sub(r'\|(?![^{}]*\})', '\n', text)
    return text.strip()
        

def compress_image(image_path, save_to, new_width_pxl: int=400):
    img_name = os.path.basename(image_path)
    try:
        print(f'compressing image {img_name}')
        with Image.open(image_path) as img:
            og_width, og_height = img.size
            new_res = (new_width_pxl, int(new_width_pxl/og_width*og_height))
            img_fix = ImageOps.exif_transpose(img)
            img_fix = img_fix.resize(new_res, Image.LANCZOS)
            img_fix.save(save_to, optimize=True, quality=85)
    except Exception as e:
        IPS()

def extract_between_substrings(text: str, delimiter: str, remaining_splitter=''):
    parts = text.split(delimiter)
    
    # If the delimiter doesn't occur in even numbers, return an empty list
    if len(parts) % 2 == 0:
        return [], text
    
    extracted = [parts[i].strip() for i in range(1, len(parts), 2)]
    remaining_raw = [parts[i] for i in range(0, len(parts), 2)]
    remainings = []
    for remaining in remaining_raw:
      if not remaining: continue
      if remaining_splitter:
        remainings.extend(remaining.strip().split(remaining_splitter))
      else: remainings.append(remaining)
    
    return [e for e in extracted if e], [r for r in remainings if r]

def Text_get(text):
    fulltext = text.get('1.0', 'end-1c')
    fulltext = fulltext.strip().replace('\n\r', '\n')
    return fulltext

def get_worst_flag(flaglist):
    flag_prio = ['PPP', 'PP']
    flag_prio.extend(list('PEIVS-*4320'))
    flag_prio.append('RAW')
    for flag in flag_prio:
        if flag in flaglist: return flag
    return ''

def timeit(f):

    def timed(*args, **kw):

        ts = time.time()
        result = f(*args, **kw)
        te = time.time()

        print(f'{np.round(te-ts, 3)} s: {f.__name__}.')
        return result

    return timed


# --- WORD FUNCTIONS ---
def simplynumber(paragraph):
    p = paragraph._p #access to xml paragraph element
    pPr = p.get_or_add_pPr() #access paragraph properties
    numPr = OxmlElement('w:numPr') #create number properties element
    numId = OxmlElement('w:numId') #create numId element - sets bullet type
    numId.set(qn('w:val'), '3') #set list type/indentation
    numPr.append(numId) #add bullet type to number properties list

    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)

    pPr.append(numPr) #add number properties to paragraph

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_TOC(run):
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Rechtsklick -> Felder aktualisieren zum Aktualisieren. Zur Entfernung der Punkte vor den Seitenzahlen: Reiter Referenzen -> Benutzerdefiniertes Inhaltsverzeichnis -> Füllzeichen: ohne, Ebenen anzeigen: 2 -> OK"
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


def autofit_table(table):
    table.autofit = True
    table.allow_autofit = True
    table._tblPr.xpath("./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
    for row_idx, r_val in enumerate(table.rows):
        for cell_idx, c_val in enumerate(table.rows[row_idx].cells):
            table.rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
            table.rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0

def set_table_column_widths(table, widths: list, relative=False, doc=None):
    '''set widths of table column in cm. length of widths mus match number of columns.'''
    if len(widths) != len(table.columns):
        raise ValueError(f'Number of widths ({len(widths)}) must match '
                         f'number of columns ({len(table.columns)})')
    if relative and (_sum:=sum(widths)) > 1: 
        raise ValueError('for relative widths, the sum of the widths must be <= 1.'
                         f'current sum: {_sum}')
    if relative:
        if not doc: raise AttributeError('When widths are relative, please pass doc.')
        sec=doc.sections[0]
        available_width = sec.page_width - (sec.left_margin + sec.right_margin)
    
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if relative: cell.width = widths[i]*available_width
            else: cell.width = Cm(widths[i])

def leftalign_table(table):
    '''change alignment within a table to left-align. Useful if table has
    linebreaks that make the text span the entire cell despite being short
    (lots of white space)'''
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_custom_footer(doc, left_text: str, middle_text: str):
    """
    Add a footer with page numbers in 'X/Y' format that appears from page 2 onward.
    
    Args:
        doc: python-docx Document object
        left_text: Text for left side of footer
        middle_text: Text for middle section of footer
    """
    # Configure section to have different first page footer
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Get regular footer (for pages after the first)
    footer = section.footer

    # Clear existing footer content
    while len(footer.paragraphs) > 0:
        delete_paragraph(footer.paragraphs[0])

    # Create paragraph and set alignment
    section = doc.sections[0]
    page_width = section.page_width.cm  # Total page width in inches
    center_pos = page_width / 2
    right_pos = page_width - section.right_margin.cm  # Account for right margin

    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(center_pos), WD_ALIGN_PARAGRAPH.CENTER)
    tab_stops.add_tab_stop(Cm(right_pos), WD_ALIGN_PARAGRAPH.RIGHT)


    # Add content with tabs
    run = paragraph.add_run()
    run.font.size = Pt(10)
    run.add_text(f"{left_text}\t{middle_text}\t")

    # Add current page number field
    _add_page_field(run)
    run.add_text("/")
    # Add total pages field
    _add_num_pages_field(run)

def _add_page_field(run):
    """Insert PAGE field into a run"""
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def _add_num_pages_field(run):
    """Insert NUMPAGES field into a run"""
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' NUMPAGES '
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Alternatively, set the run's formatting explicitly
    new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    new_run.font.underline = False

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def open_file(path):
    if sys.platform.startswith('darwin'):  # macOS
        subprocess.call(('open', path))
    elif os.name == 'nt':  # Windows
        os.startfile(path)
    elif os.name == 'posix':  # Linux, BSD, etc.
        subprocess.call(('xdg-open', path))


    
